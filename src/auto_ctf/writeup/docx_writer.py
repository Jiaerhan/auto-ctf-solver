""".docx writeup writer using python-docx."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


class DocxWriter:
    """Writes CTF writeups as Word (.docx) documents with embedded screenshots."""

    def write(self, data: dict, output_path: Path) -> Path:
        doc = Document()

        # Title
        title = doc.add_heading(f"CTF Writeup: {data['name']}", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Metadata
        meta = doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta.add_run(f"Category: {data['category']}").bold = True
        meta.add_run(f"\nDate: {data['timestamp'][:10]}")
        if data.get("duration"):
            meta.add_run(f"\nDuration: {data['duration']:.1f}s")

        doc.add_paragraph()

        # Challenge Info
        doc.add_heading("Challenge Information", level=1)
        doc.add_paragraph(f"Category: {data['category']}")
        if data.get("flag"):
            doc.add_paragraph(f"Flag: {data['flag']}")

        # Approach
        doc.add_heading("Solution Approach", level=1)
        doc.add_paragraph(data.get("approach", "No approach recorded."))

        # Steps
        doc.add_heading("Solution Steps", level=1)
        for i, step in enumerate(data.get("steps", []), 1):
            doc.add_heading(f"Step {i}", level=2)
            doc.add_paragraph(step[:2000])  # Truncate very long steps

        # Flag
        if data.get("flag"):
            doc.add_heading("Flag", level=1)
            p = doc.add_paragraph()
            run = p.add_run(data["flag"])
            run.bold = True
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)

        # Screenshots
        screenshots = data.get("screenshots", [])
        if screenshots:
            doc.add_heading("Screenshots", level=1)
            for shot in screenshots:
                shot_path = Path(shot)
                if shot_path.exists():
                    try:
                        doc.add_picture(str(shot_path), width=Inches(5.5))
                        doc.add_paragraph(shot_path.name).italic = True
                    except Exception:
                        doc.add_paragraph(f"[Screenshot: {shot}]")

        # Errors (if any)
        if data.get("errors"):
            doc.add_heading("Errors / Notes", level=1)
            for err in data["errors"]:
                doc.add_paragraph(err, style="List Bullet")

        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))
        return output_path
